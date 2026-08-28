import re
import os
import urllib.request
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from lxml import etree
import latex2mathml.converter

# =============================================================================
# Dependencies Setup:
# pip install python-docx latex2mathml lxml
# =============================================================================

# We need Microsoft's MathML to OMML XSLT stylesheet. 
XSLT_URL = "https://raw.githubusercontent.com/transpect/docx_modify-lib/master/xsl/mml2omml.xsl"
XSLT_FILE = "MML2OMML.XSL"

def ensure_xslt_exists():
    """Downloads the MML2OMML stylesheet if not present."""
    if not os.path.exists(XSLT_FILE):
        print(f"Downloading {XSLT_FILE}...")
        urllib.request.urlretrieve(XSLT_URL, XSLT_FILE)
        print("Download complete.")

def latex_to_omml(latex_str, xslt_transform):
    """
    Converts a LaTeX math string into a Word OMML XML element.
    """
    # 1. Convert LaTeX to MathML
    mathml_str = latex2mathml.converter.latex2mathml(latex_str)
    
    # 2. Parse MathML into lxml element
    mathml_xml = etree.fromstring(mathml_str)
    
    # 3. Apply XSLT to convert MathML to OMML
    omml_xml = xslt_transform(mathml_xml)
    
    # 4. Convert lxml element back to string, then to docx oxml element
    omml_str = str(omml_xml)
    
    # Clean up the OMML string for python-docx
    omml_str = omml_str.replace('<?xml version="1.0"?>', '').strip()
    
    # Ensure correct namespace for OMML
    if not 'xmlns:m=' in omml_str:
        omml_str = omml_str.replace('<m:oMath>', f'<m:oMath {nsdecls("m")}>')
    
    return parse_xml(omml_str)

def process_text_to_docx(input_text, output_docx_path):
    """
    Parses plain text with LaTeX equations and generates a Word document.
    """
    ensure_xslt_exists()
    
    # Load XSLT Transformer
    xslt_doc = etree.parse(XSLT_FILE)
    xslt_transform = etree.XSLT(xslt_doc)
    
    doc = Document()
    
    # Regex to find $$...$$ or $...$
    # It captures the delimiters and the math content
    pattern = re.compile(r'(\$\$.*?\$\$|\$.*?\$)')
    
    # Process paragraph by paragraph (split by newlines)
    paragraphs = input_text.split('\n')
    
    for text_para in paragraphs:
        if not text_para.strip():
            continue
            
        p = doc.add_paragraph()
        
        # Split paragraph into text and math parts
        parts = pattern.split(text_para)
        
        for part in parts:
            if part.startswith('$$') and part.endswith('$$'):
                # Display Math
                latex_code = part[2:-2].strip()
                try:
                    omml_elem = latex_to_omml(latex_code, xslt_transform)
                    p._element.append(omml_elem)
                except Exception as e:
                    print(f"Failed to convert: {latex_code} - {e}")
                    p.add_run(part) # Fallback to plain text
                    
            elif part.startswith('$') and part.endswith('$'):
                # Inline Math
                latex_code = part[1:-1].strip()
                try:
                    omml_elem = latex_to_omml(latex_code, xslt_transform)
                    p._element.append(omml_elem)
                except Exception as e:
                    print(f"Failed to convert: {latex_code} - {e}")
                    p.add_run(part) # Fallback to plain text
                    
            else:
                # Regular Text
                if part:
                    p.add_run(part)
                    
    doc.save(output_docx_path)
    print(f"Successfully saved Word document with native math: {output_docx_path}")

if __name__ == "__main__":
    # Example usage:
    sample_text = """
    Here is an example of an inline equation: $E=mc^2$.
    
    And here is a display equation (quadratic formula):
    $$x = \frac{-b \pm \sqrt{b^2 - 4ac}}{2a}$$
    
    Word will render these perfectly inline with the text!
    """
    
    output_file = "Math_Output.docx"
    process_text_to_docx(sample_text, output_file)
